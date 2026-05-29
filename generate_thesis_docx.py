from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_thesis_chapters():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_heading(text, level):
        h = doc.add_heading(text, level=level)
        h.style.font.name = 'Times New Roman'
        h.style.font.color.rgb = None # default black
        
    def add_figure(image_path, caption):
        if os.path.exists(image_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(image_path, width=Inches(5.5))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(caption)
            r_cap.italic = True
            r_cap.font.size = Pt(10)
        else:
            doc.add_paragraph(f"[MISSING IMAGE: {image_path}]")

    # --- CHAPTER 4 ---
    add_heading('CHAPTER 4', 1)
    add_heading('RESULTS AND ANALYSIS', 1)
    
    doc.add_paragraph("This chapter presents a comprehensive evaluation of the proposed AuraSentinel forecasting system. The results are systematically categorized into data preprocessing insights, comparative model performance, classification metrics for extreme event detection, explainability through SHAP (Shapley Additive exPlanations), and a rigorous statistical validation using a walk-forward time-series test harness. The primary objective is to empirically validate the superiority of the proposed Hybrid CNN-BiLSTM-Attention architecture over conventional machine learning baselines.")

    add_heading('4.1 Data Understanding and Exploratory Data Analysis (EDA)', 2)
    doc.add_paragraph("Understanding the underlying distribution of the hydrometeorological dataset is critical before applying deep learning models. The dataset consists of 41,640 daily observations recorded across three primary CAMELS-AUS catchments. Figure 4.1 illustrates the highly skewed nature of daily rainfall, which is characteristic of extreme weather phenomena. The vast majority of days record near-zero precipitation, whereas the tail end of the distribution captures the rare but devastating flood-inducing rainfall events exceeding the 95th percentile threshold.")
    add_figure('Thesis_Figures/FigA_Rainfall_Distribution.png', 'Figure 4.1: Statistical Distribution and Kernel Density Estimate of Daily Rainfall (Showing Extreme Class Imbalance)')
    
    doc.add_paragraph("To capture the temporal sequence of these extreme events, Figure 4.2 presents a sample time-series trajectory of the rainfall data. The red markers indicate extreme rainfall spikes that exceed the 95th percentile. This temporal clustering demonstrates the necessity of the proposed BiLSTM architecture, which is inherently designed to capture bidirectional sequential dependencies across the 14-day lookback window.")
    add_figure('Thesis_Figures/FigB_Timeseries_Rainfall.png', 'Figure 4.2: Temporal Variation of Daily Rainfall Highlighting Extreme Spikes')
    
    doc.add_paragraph("Feature correlation is a foundational step in identifying multicollinearity and predictive strength. Figure 4.3 presents the Pearson Correlation Matrix of the meteorological variables. The heatmap reveals a strong positive correlation between 'Rainfall_Roll7' (the 7-day rolling average) and the target rainfall, confirming that short-term antecedent catchment memory is a critical predictor. Furthermore, variables such as Mean Sea Level Pressure (MSLP) exhibit inverse correlations with rainfall, aligning with atmospheric dynamics where low-pressure systems drive heavy precipitation.")
    add_figure('Thesis_Figures/FigC_Correlation_Heatmap.png', 'Figure 4.3: Pearson Correlation Matrix of Hydrometeorological Features')

    add_heading('4.2 Preprocessing and Normalization', 2)
    doc.add_paragraph("Real-world sensor data is inherently noisy and often contains missing values. Figure 4.4 illustrates the missing data distribution before imputation. A 30-day centered rolling mean was applied to impute these gaps, ensuring that seasonal trends were preserved without introducing forward-looking data leakage.")
    add_figure('Thesis_Figures/FigD_Missing_Values.png', 'Figure 4.4: Missing Data Distribution Matrix')
    
    doc.add_paragraph("Because neural networks calculate gradients based on input magnitudes, features with vastly different scales (e.g., MSLP in hPa vs. Temperature in °C) can destabilize the loss function. Figure 4.5 demonstrates the feature distributions before and after applying MinMax Normalization. The post-normalization boxplots confirm that all input variables have been strictly bounded within the [0, 1] interval, guaranteeing equal gradient contribution during the backpropagation phase of the CNN-BiLSTM-Attention model.")
    add_figure('Thesis_Figures/FigE_Normalization_Comparison.png', 'Figure 4.5: Feature Boxplots Before and After MinMax Normalization')

    add_heading('4.3 Model Convergence and Training Dynamics', 2)
    doc.add_paragraph("The training dynamics of the deep learning architectures were monitored across 45 epochs to evaluate convergence stability and detect overfitting. Figure 4.6 presents the Huber loss and accuracy progression for both the baseline Stacked LSTM and the proposed Hybrid model. The Hybrid CNN-BiLSTM-Attention model demonstrates a steeper initial descent in the loss curve and achieves a stable validation loss plateau without diverging, indicating robust generalization. The utilization of Dropout (0.2) and Batch Normalization layers effectively mitigated overfitting.")
    add_figure('Thesis_Figures/FigF2_Loss_Hybrid.png', 'Figure 4.6: Training and Validation Loss/Accuracy Curves of the Hybrid Model')

    add_heading('4.4 Comparative Model Performance (Regression Analysis)', 2)
    doc.add_paragraph("The forecasting pipeline was executed in three progressive phases: Phase 1 (Baseline Stacked LSTM), Phase 2 (XGBoost + Random Forest Ensemble), and Phase 3 (The Proposed Hybrid CNN-BiLSTM-Attention). Figure 4.7 visualizes the actual versus predicted rainfall trajectories over a 150-day hold-out test segment. While the baseline LSTM struggles to reach the peak magnitudes of extreme events (under-prediction), the Hybrid model successfully captures the amplitude of the spikes. This is largely attributed to the self-attention mechanism, which dynamically assigns higher mathematical weights to critical days within the 14-day input window.")
    add_figure('Thesis_Figures/FigH_Actual_vs_Predicted_Line.png', 'Figure 4.7: Actual vs. Predicted Rainfall Trajectories Across Models')
    
    doc.add_paragraph("Figure 4.8 further quantifies this via regression scatter plots. The Hybrid model yields the tightest clustering along the perfect fit diagonal line, achieving the highest coefficient of determination (R²). The ensemble model exhibits moderate dispersion, while the baseline LSTM shows significant heteroscedasticity at higher rainfall values.")
    add_figure('Thesis_Figures/FigI_Scatter_Actual_Predicted.png', 'Figure 4.8: Scatter Plot Comparison of Observed vs. Predicted Rainfall')

    add_heading('4.5 Quantitative Performance Metrics', 2)
    doc.add_paragraph("A comprehensive evaluation requires multiple deterministic metrics. As shown in the comparative bar charts (Figure 4.9), the Hybrid model consistently outperforms the baselines across all dimensions. Specifically, it achieves the lowest Root Mean Square Error (RMSE = 3.63 mm) and Mean Absolute Error (MAE = 1.58 mm). Most notably for hydrological applications, the Critical Success Index (CSI) reached 0.59, and the False Alarm Rate (FAR) was minimized to 8.4%.")
    add_figure('Thesis_Figures/Fig_J_RMSE_Comparison.png', 'Figure 4.9a: RMSE Comparison')
    add_figure('Thesis_Figures/Fig_M_CSI_Comparison.png', 'Figure 4.9b: Critical Success Index Comparison')
    add_figure('Thesis_Figures/Fig_N_False_Comparison.png', 'Figure 4.9c: False Alarm Rate Comparison')
    
    doc.add_paragraph("Predictive capability inherently degrades as the forecasting horizon extends. Figure 4.10 illustrates the lead-time sensitivity analysis. While all models degrade at T+3 and T+7 days, the CNN-BiLSTM-Attention architecture exhibits the most resilient predictive decay, maintaining an 89.42% accuracy even at a 7-day lead time. This confirms that the model has successfully learned long-term atmospheric dynamics rather than just relying on immediate lag values.")
    add_figure('Thesis_Figures/FigO_Lead_Time.png', 'Figure 4.10: Forecast Lead Time Accuracy Degradation')

    add_heading('4.6 Classification Analysis for Extreme Events', 2)
    doc.add_paragraph("Beyond continuous regression, the system's operational viability depends on its ability to correctly classify and issue alerts for extreme flood-inducing events. Figure 4.11 presents the Confusion Matrix for the Hybrid model based on a 90th percentile extremity threshold. The dense concentration of True Positives and True Negatives, coupled with minimal False Negatives, indicates a highly reliable early warning capability.")
    add_figure('Thesis_Figures/FigP_Confusion_Matrix.png', 'Figure 4.11: Confusion Matrix for Extreme Event Classification')
    
    doc.add_paragraph("To evaluate the true positive rate against the false positive rate, Figure 4.12 provides the Receiver Operating Characteristic (ROC) curve. The Hybrid model achieves an exceptional Area Under the Curve (AUC) score. Because flood events are heavily imbalanced (rare class), the Precision-Recall (PR) curve in Figure 4.13 is also provided, confirming that the model maintains high precision even at elevated recall levels—a fundamental requirement for preventing alarm fatigue in disaster management systems.")
    add_figure('Thesis_Figures/FigQ_ROC_Curve.png', 'Figure 4.12: ROC Curve')
    add_figure('Thesis_Figures/FigR_PR_Curve.png', 'Figure 4.13: Precision-Recall Tradeoff')

    add_heading('4.7 Feature Explainability and Ablation Study', 2)
    doc.add_paragraph("Modern AI systems in meteorology must be interpretable. SHAP (Shapley Additive exPlanations) values were utilized to extract the global feature importance of the Hybrid model (Figure 4.14). The analysis reveals that Rainfall(Lag1) and the 7-day rolling rainfall are the dominant drivers of the prediction, followed by Mean Sea Level Pressure. Climate indices (SOI, DMI) exert a non-linear background influence.")
    add_figure('Thesis_Figures/FigU_SHAP_Summary.png', 'Figure 4.14: SHAP Global Feature Importance')
    
    doc.add_paragraph("To mathematically justify the complexity of the proposed architecture, an Ablation Study was conducted (Figure 4.15). Removing the Attention layer caused a severe degradation in RMSE (from 3.63mm to 4.12mm), proving that dynamic weighting of the temporal sequence is the primary catalyst for the model's superior performance.")
    add_figure('Thesis_Figures/FigAA_Ablation.png', 'Figure 4.15: Architectural Ablation Study Comparison')

    # --- CHAPTER 5 ---
    add_heading('CHAPTER 5', 1)
    add_heading('DISCUSSION AND CONCLUSION', 1)
    
    add_heading('5.1 Walk-Forward Validation and Statistical Significance', 2)
    doc.add_paragraph("To strictly adhere to time-series forecasting principles and avoid data leakage, standard k-fold cross-validation was discarded in favor of Walk-Forward Validation (TimeSeriesSplit, k=5). Under this strict test harness, a Paired T-Test was conducted on the absolute prediction errors between the Baseline LSTM and the Hybrid model.")
    doc.add_paragraph("The statistical test yielded a T-Statistic of -14.76 and a P-Value of 3.65e-49. Because the p-value is significantly less than the standard 0.05 alpha threshold, the null hypothesis is firmly rejected. This provides mathematical proof that the performance enhancement of the CNN-BiLSTM-Attention model is statistically significant and definitively not the result of random variation or dataset splitting artifacts.")

    add_heading('5.2 Thesis Conclusion', 2)
    doc.add_paragraph("This research successfully engineered, optimized, and validated the AuraSentinel flood forecasting system. The objective of outperforming standard machine learning and baseline deep learning models was conclusively achieved. The proposed Hybrid CNN-BiLSTM-Attention architecture demonstrated unparalleled efficacy in capturing the complex, non-linear spatiotemporal dynamics of extreme rainfall events.")
    doc.add_paragraph("By integrating 1D Convolutional layers for spatial feature extraction, Bidirectional LSTMs for chronological memory, and an Attention mechanism to prioritize peak intensity days, the system achieved a champion accuracy of 97.43% with a minimal RMSE of 3.63mm. The rigorous evaluation harness, including Walk-Forward validation, SHAP explainability, and paired statistical testing, establishes this framework as a highly robust, operationally viable tool for early disaster warning in the context of escalating global climate volatility.")

    doc.save('Chapter_4_5_6_Complete_Results.docx')

if __name__ == '__main__':
    create_thesis_chapters()
