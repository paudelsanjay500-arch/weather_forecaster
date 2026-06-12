from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

def generate_raw_materials_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    doc.add_heading('Raw Results and Figures for Claude', level=1)
    doc.add_paragraph("This document contains the genuine evaluation metrics and generated figures. Please use this data to write the comprehensive Chapters 4, 5, and 6.")
    
    # 1. Add Metrics Tables
    doc.add_heading('1. Genuine Evaluation Metrics', level=2)
    
    json_files = ['phase1_results.json', 'ensemble_metrics.json', 'hybrid_metrics.json']
    
    for j_file in json_files:
        if os.path.exists(j_file):
            with open(j_file, 'r') as f:
                data = json.load(f)
            
            doc.add_heading(data.get('Phase', j_file), level=3)
            
            # Create a table for the metrics
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            for key, value in data.items():
                if key != 'Phase':
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
                    
            doc.add_paragraph("\n")
            
    # 2. Add Figures
    doc.add_heading('2. Generated Figures with Captions', level=2)
    
    figures = [
        ('Thesis_Figures/Fig4_1_Correlation_Heatmap.png', 'Figure 4.1: Pearson Correlation Matrix of Hydrometeorological Features'),
        ('Thesis_Figures/Fig4_2_Rainfall_Distribution.png', 'Figure 4.2: Statistical Distribution and Kernel Density Estimate of Daily Rainfall'),
        ('Thesis_Figures/Fig4_3_Monthly_Seasonality.png', 'Figure 4.3: Seasonal Variation and Monthly Patterns of Rainfall'),
        ('Thesis_Figures/Fig4_4_Scatter_Regression.png', 'Figure 4.4: Scatter Plot Comparison of Observed vs. Predicted Rainfall'),
        ('Thesis_Figures/Fig4_5_Confusion_Matrices.png', 'Figure 4.5: Confusion Matrices for Extreme Event Classification Across Models'),
        ('Thesis_Figures/Fig4_6_WalkForward_RMSE.png', 'Figure 4.6: Walk-Forward Validation RMSE Across Time Splits'),
        ('Thesis_Figures/Fig4_7_Model_Comparison_Bar.png', 'Figure 4.7: Quantitative Performance Metrics Comparison (Accuracy, RMSE, NSE, FAR)'),
        ('Thesis_Figures/Fig4_8_LeadTime_Sensitivity.png', 'Figure 4.8: Forecast Lead Time Accuracy Degradation (T+1 to T+7 Days)'),
        ('Thesis_Figures/Fig4_9_Residual_Analysis.png', 'Figure 4.9: Residual Analysis for Model Prediction Errors'),
        ('Thesis_Figures/Fig4_10_Feature_Importance.png', 'Figure 4.10: SHAP Global Feature Importance and Ablation Analysis'),
        ('Thesis_Figures/Fig4_11_ROC_Curve.png', 'Figure 4.11: Receiver Operating Characteristic (ROC) Curve'),
        ('Thesis_Figures/Fig4_12_PR_Curve.png', 'Figure 4.12: Precision-Recall (PR) Curve for Imbalanced Extreme Events'),
        ('Thesis_Figures/Fig4_13_Loss_Curves.png', 'Figure 4.13: Training vs. Validation Loss Convergence')
    ]
    
    for img_path, caption in figures:
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(img_path, width=Inches(6.0))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(caption)
            r_cap.italic = True
            r_cap.font.size = Pt(10)
        else:
            doc.add_paragraph(f"[ERROR: Missing Image {img_path}]")
        
    output_name = 'Raw_Results_Materials_For_Claude_v3.docx'
    doc.save(output_name)
    print(f"Successfully generated {output_name}")

if __name__ == '__main__':
    generate_raw_materials_docx()
